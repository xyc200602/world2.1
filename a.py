import graphviz
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ==========================================
# 1. 生成 Word 文档 (Detailed HTA Report)
# ==========================================
def create_hta_document(filename):
    doc = Document()
    
    # 标题
    heading = doc.add_heading('基于 INSARAG 标准的单兵城市救援任务创新型 HTA', 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 简介与符号说明
    doc.add_heading('一、 符号与定义说明 (Innovation Framework)', level=1)
    intro = doc.add_paragraph()
    intro.add_run('本分析报告基于《INSARAG Guidelines Vol II Manual B》构建，采用 E-HTA (Extended Hierarchical Task Analysis) 框架。').italic = True
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '符号'
    hdr_cells[1].text = '定义与内涵'
    
    definitions = [
        ("G_dyn (Dynamic Goal)", "动态目标：包含优先级权重与中止阈值的战略目标。"),
        ("P_trig (Trigger Plan)", "触发逻辑：基于事件驱动（Event-Driven）的控制流，负责中断或切换流程。"),
        ("T (Task)", "操作：具体的物理执行动作。"),
        ("E_c (Env. Constraint)", "环境约束：现场物理环境对感官或行动的限制。"),
        ("I_r (Info. Requirement)", "信息需求：执行操作前必须获取的关键数据输入。"),
        ("CP (Collab. Protocol)", "协同协议：与其他人员或系统的交互规则。")
    ]
    
    for symbol, desc in definitions:
        row_cells = table.add_row().cells
        row_cells[0].text = symbol
        row_cells[1].text = desc

    # HTA 正文内容生成函数
    def add_hta_module(title, g_dyn, p_trig, operations):
        doc.add_heading(title, level=1)
        
        # 动态目标
        p = doc.add_paragraph()
        run = p.add_run(f"G_dyn (动态目标): {g_dyn}")
        run.bold = True
        run.font.color.rgb = RGBColor(0, 51, 102) # Dark Blue
        
        # 触发逻辑
        p = doc.add_paragraph()
        run = p.add_run(f"P_trig (触发逻辑): {p_trig}")
        run.bold = True
        run.font.color.rgb = RGBColor(153, 0, 0) # Dark Red
        
        # 操作详情
        for op_title, details in operations.items():
            doc.add_heading(op_title, level=2)
            for key, value in details.items():
                p = doc.add_paragraph(style='List Bullet')
                key_run = p.add_run(f"{key}: ")
                key_run.bold = True
                p.add_run(value)

    # --- 模块 1 ---
    add_hta_module(
        "模块 1: 现场分检与评估 (Worksite Triage - ASR2)",
        "基于存活率计算救援优先级。若确认幸存者且耗时<12小时，权重置顶 (Cat A)；若仅遇难者，权重最低 (Cat D) 。",
        "IF 发现无法隔离的危险品 (Hazmat) -> THEN 中止分检并报告 ；IF 幸存者位置确认 -> THEN 跳转至标记模块。",
        {
            "1.1 操作: 搜集结构特征": {
                "T (Task)": "观察倒塌模式（倾斜/层叠/废墟堆）及空隙类型 。",
                "E_c (环境约束)": "E_visual: 视线受阻；E_debris: 废墟堆导致无法观察底部。",
                "I_r (信息需求)": "建筑用途（学校/医院？）；构造材料（重型/轻型？）。"
            },
            "1.2 操作: 获取受困者情报": {
                "T (Task)": "整合多源情报验证生命迹象 。",
                "CP (协同协议)": "CP_LEMA: 必须与当地机构核对失踪报告；CP_Local: 询问旁观者区分'可能'与'确认' 。",
                "I_r (信息需求)": "判定逻辑：USAR确认='Confirmed'；旁观者报告='Possible'。"
            }
        }
    )

    # --- 模块 2 ---
    add_hta_module(
        "模块 2: 结构标记与通信 (Marking System)",
        "建立异步持久化信息节点。中止阈值：结构极度不稳定导致无法靠近主入口 。",
        "IF ASR2 完成 -> THEN 绘制ID框；IF 发现危险品 -> THEN 在框外标注明文 ；IF 救援结束 -> THEN 画水平线划掉 。",
        {
            "2.1 操作: 绘制工作面 ID": {
                "T (Task)": "在入口绘制 1.2m x 1.0m 方框及 40cm ID 。",
                "E_c (环境约束)": "E_surface: 废墟表面粗糙需强附着力材料；E_contrast: 需高对比度颜色 。",
                "I_r (信息需求)": "必需字段：队伍代码、ASR等级、日期 ；动态字段：后续队伍追加记录。"
            },
            "2.2 操作: 受困者定位标记": {
                "T (Task)": "在受困点附近喷涂 'V'，下方标注 L (活) 或 D (死) 。",
                "CP (协同协议)": "CP_Async: 给后续队伍看，需画在物理最近点而非门口 。",
                "I_r (信息需求)": "更新逻辑：救出一人，划掉 L-2 改写 L-1 。"
            }
        }
    )

    # --- 模块 3 ---
    add_hta_module(
        "模块 3: 搜救执行 (Operations - ASR3/4)",
        "最大化救出率。优先级翻转：若受困太深，ASR3 目标降级为'标记并移交'，除非指令升级 ASR4 。",
        "IF 耗时 > 1作业周期 -> THEN 中止 ASR3 ；IF 发现深层受困者 -> THEN 禁止深入挖掘 。",
        {
            "3.1 操作: 浅层搜救 (ASR3)": {
                "T (Task)": "移除表面废墟，有限支撑，不深入结构内部 。",
                "E_c (环境约束)": "E_time: 仅有数小时时间窗；E_access: 仅限表面空隙。",
                "I_r (信息需求)": "是否具备快速移除的条件？"
            },
            "3.2 操作: 深层重型搜救 (ASR4)": {
                "T (Task)": "切割重型元件，建立深层通道，全面支撑 。",
                "CP (协同协议)": "CP_Sync: 需现场完全指挥控制；CP_Logistics: 配合重型机械 。",
                "I_r (信息需求)": "结构应力分析：移除梁是否导致二次坍塌？"
            }
        }
    )
    
    # --- 模块 4 ---
    add_hta_module(
        "模块 4: 危险响应 (Safety & Signals)",
        "保障自身存活 (Self-Preservation)。权重：Override 所有其他目标。",
        "IF 听到信号 -> IMMEDIATELY 触发反射动作 ；IF 监测到 Hazmat -> THEN 建立警戒线并撤离 。",
        {
            "4.1 操作: 紧急撤离": {
                "T (Task)": "丢弃装备，沿路线撤离。",
                "I_r (信息需求)": "信号特征：3次短促信号 (1秒/次) 。",
                "CP (协同协议)": "CP_Universal: 全员统一反应 。"
            },
            "4.2 操作: 静默 (声波探测)": {
                "T (Task)": "停止动作，关闭引擎 。",
                "I_r (信息需求)": "信号特征：1次长信号 (3秒) 。",
                "E_c (环境约束)": "E_noise: 现场极其嘈杂，需汽笛覆盖背景音 。"
            }
        }
    )

    doc.save(filename)
    return filename

# ==========================================
# 执行部分 - 添加这里！
# ==========================================
if __name__ == "__main__":
    try:
        # 生成Word文档
        doc_filename = "INSARAG_HTA_Report.docx"
        result = create_hta_document(doc_filename)
        print(f"✅ Word文档生成成功: {result}")
        
        # 检查文件是否存在
        import os
        if os.path.exists(doc_filename):
            file_size = os.path.getsize(doc_filename)
            print(f"📁 文件大小: {file_size} 字节")
        else:
            print("❌ 文件未生成，请检查权限或路径")
            
    except Exception as e:
        print(f"❌ 生成过程中出现错误: {str(e)}")